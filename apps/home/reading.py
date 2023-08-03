from docx import Document

def read(docx_file, paragraph_index):
    doc = Document(docx_file)

    # Ensure paragraph_index is within valid range
    if 0 <= paragraph_index < len(doc.paragraphs):
        return doc.paragraphs[paragraph_index].text
    else:
        return "Paragraph index out of range."

# def read_paragraphs_range(docx_file, start_index, end_index):
#     doc = Document(docx_file)

#     # Ensure start and end indexes are within valid range
#     if 0 <= start_index < len(doc.paragraphs) and 0 <= end_index < len(doc.paragraphs):
#         return '\n'.join(doc.paragraphs[start_index:end_index+1])
#     else:
#         return "Start or end index out of range."

# # Usage examples:
# docx_file = 'input.docx'  # Replace with the path to your .docx file

# # Read specific paragraph
# paragraph_index = 2  # Replace with the desired paragraph index
# specific_paragraph = read_specific_paragraph(docx_file, paragraph_index)
# print(specific_paragraph)

# # Read a range of paragraphs
# start_index = 2  # Replace with the start paragraph index
# end_index = 5    # Replace with the end paragraph index
# paragraphs_range = read_paragraphs_range(docx_file, start_index, end_index)
# print(paragraphs_range)
